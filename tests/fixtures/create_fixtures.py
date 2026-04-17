"""
Helper script to generate binary fixture files (sample_cv.docx, sample_cv.pdf).
Run once from the project root: python tests/fixtures/create_fixtures.py
"""

import os

FIXTURES_DIR = os.path.dirname(os.path.abspath(__file__))


def create_docx():
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Contact header
    doc.add_heading("Jane Doe", level=0)
    doc.add_paragraph("jane.doe@email.com | +44 7700 900123 | linkedin.com/in/janedoe")

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Results-driven software engineer with 7 years of experience designing and "
        "delivering scalable backend systems. Proven track record of leading cross-functional "
        "teams, improving system reliability, and shipping high-quality Python and cloud-native "
        "solutions. Passionate about clean architecture, automated testing, and continuous delivery."
    )

    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(
        "Python, Java, SQL, PostgreSQL, MySQL, AWS, Docker, Kubernetes, Terraform, CI/CD, "
        "Jenkins, GitHub Actions, REST APIs, Microservices, Agile, Scrum, Git, Linux, Redis, Elasticsearch"
    )

    # Work Experience
    doc.add_heading("Work Experience", level=1)

    doc.add_heading("Senior Software Engineer", level=2)
    doc.add_paragraph("Acme Technologies Ltd, London, UK")
    doc.add_paragraph("03/2021 – Present")
    bullets_1 = [
        "Designed and implemented a microservices-based order processing platform handling "
        "50,000 transactions per day using Python and FastAPI.",
        "Led migration of monolithic application to AWS ECS, reducing infrastructure costs by 35% "
        "and improving deployment frequency from monthly to daily.",
        "Introduced automated integration testing with pytest, increasing code coverage from 42% to 87%.",
        "Mentored a team of 4 junior engineers, conducting weekly code reviews and pair programming sessions.",
        "Collaborated with product managers and stakeholders to define technical requirements and delivery roadmaps.",
    ]
    for bullet in bullets_1:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Software Engineer", level=2)
    doc.add_paragraph("Brightwave Solutions, Manchester, UK")
    doc.add_paragraph("06/2018 – 02/2021")
    bullets_2 = [
        "Built and maintained RESTful APIs consumed by web and mobile clients, serving over 200,000 active users.",
        "Implemented a real-time notification service using Redis Pub/Sub and WebSockets, reducing notification latency by 60%.",
        "Optimised slow PostgreSQL queries, cutting average API response time from 800 ms to 120 ms.",
        "Contributed to the adoption of Docker and Kubernetes for local development and staging environments.",
        "Participated in on-call rotation, resolving production incidents and writing post-mortems.",
    ]
    for bullet in bullets_2:
        doc.add_paragraph(bullet, style="List Bullet")

    # Education
    doc.add_heading("Education", level=1)
    doc.add_heading("Bachelor of Science in Computer Science", level=2)
    doc.add_paragraph("University of Manchester, Manchester, UK")
    doc.add_paragraph("09/2014 – 06/2018")
    doc.add_paragraph("First Class Honours")

    # Certifications
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("AWS Certified Solutions Architect – Associate, Amazon Web Services, 2022")
    doc.add_paragraph("Certified Kubernetes Administrator (CKA), Cloud Native Computing Foundation, 2023")

    out_path = os.path.join(FIXTURES_DIR, "sample_cv.docx")
    doc.save(out_path)
    print(f"Created: {out_path}")


def create_pdf_placeholder():
    """
    Creates a placeholder text file explaining how to generate the PDF manually.
    A real PDF can be generated with reportlab or by exporting sample_cv.docx via LibreOffice:
        libreoffice --headless --convert-to pdf tests/fixtures/sample_cv.docx --outdir tests/fixtures/
    """
    placeholder_path = os.path.join(FIXTURES_DIR, "sample_cv.pdf.txt")
    content = """\
HOW TO GENERATE sample_cv.pdf
==============================

The sample_cv.pdf fixture is a PDF version of sample_cv.txt / sample_cv.docx.

Option 1 — LibreOffice (recommended, no extra Python deps):
    libreoffice --headless --convert-to pdf tests/fixtures/sample_cv.docx --outdir tests/fixtures/

Option 2 — reportlab (Python):
    pip install reportlab
    Then run the snippet below:

    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet

    doc = SimpleDocTemplate("tests/fixtures/sample_cv.pdf", pagesize=A4)
    styles = getSampleStyleSheet()
    content_lines = open("tests/fixtures/sample_cv.txt").read().splitlines()
    story = [Paragraph(line or "&nbsp;", styles["Normal"]) for line in content_lines]
    doc.build(story)

The PDF content should match sample_cv.txt exactly so that parse_cv tests
can verify round-trip fidelity across all three supported formats.
"""
    with open(placeholder_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Created: {placeholder_path}")


if __name__ == "__main__":
    create_docx()
    create_pdf_placeholder()
    print("Done.")
