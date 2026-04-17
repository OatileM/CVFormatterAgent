"""Tool: generate_output

Writes the reformatted CV and keyword report to .txt and .docx files.
Reads all data from session state — no large arguments needed.
"""

from __future__ import annotations

import os

from docx import Document
from strands import tool

import session
from models import KeywordReport, OutputPaths, ReformattedCV, ToolError

_SEPARATOR = "=" * 80


@tool
def generate_output(output_dir: str = "output") -> str:
    """Write the reformatted CV and keyword report to .txt and .docx files.

    Call this last, after reformat_cv has completed.

    Args:
        output_dir: Directory to write output files into (default: "output").

    Returns:
        A short status message with the output file paths, or an error description.
    """
    try:
        if session.reformatted_cv is None:
            return ToolError(
                tool_name="generate_output",
                error_description="No reformatted CV found. Call reformat_cv first.",
                suggested_action="Call reformat_cv before generate_output.",
            ).to_json()

        if session.keyword_report is None:
            return ToolError(
                tool_name="generate_output",
                error_description="No keyword report found. Call analyze_keywords first.",
                suggested_action="Call analyze_keywords before generate_output.",
            ).to_json()

        os.makedirs(output_dir, exist_ok=True)

        txt_path = os.path.join(output_dir, "cv.txt")
        docx_path = os.path.join(output_dir, "cv.docx")

        report_block = _build_report_block(session.keyword_report)

        # Write .txt
        full_text = session.reformatted_cv.plain_text
        if not full_text.endswith("\n"):
            full_text += "\n"
        full_text += "\n" + report_block
        with open(txt_path, "w", encoding="utf-8") as fh:
            fh.write(full_text)

        # Write .docx
        _write_docx(session.reformatted_cv, report_block, docx_path)

        return (
            f"Output files written successfully.\n"
            f"  TXT:  {txt_path}\n"
            f"  DOCX: {docx_path}\n"
            f"  ATS Score: {session.keyword_report.ats_score}%"
        )

    except Exception as exc:
        return ToolError(
            tool_name="generate_output",
            error_description=f"Failed to generate output files: {exc}",
            suggested_action="Ensure the output directory is writable and retry.",
        ).to_json()


def _build_report_block(report: KeywordReport) -> str:
    lines = [
        _SEPARATOR,
        "ATS KEYWORD MATCH REPORT",
        _SEPARATOR,
        f"ATS Score: {report.ats_score}%",
        "",
        f"Required Keywords Present ({len(report.present_required)}):",
        *[f"  - {kw}" for kw in report.present_required],
        "",
        f"Required Keywords Absent ({len(report.absent_required)}):",
        *[f"  - {kw}" for kw in report.absent_required],
        "",
        f"Preferred Keywords Present ({len(report.present_preferred)}):",
        *[f"  - {kw}" for kw in report.present_preferred],
        "",
        f"Preferred Keywords Absent ({len(report.absent_preferred)}):",
        *[f"  - {kw}" for kw in report.absent_preferred],
    ]

    if report.ats_score < 70:
        lines += [
            "",
            "RECOMMENDATIONS",
            "-" * 15,
            "Your ATS score is below 70. Consider addressing these absent required keywords:",
            *[f"  - {kw}: Add relevant experience or skills that demonstrate this competency."
              for kw in report.absent_required],
        ]

    return "\n".join(lines)


def _write_docx(reformatted_cv: ReformattedCV, report_block: str, docx_path: str) -> None:
    doc = Document()
    for section in reformatted_cv.sections:
        if section.heading:
            doc.add_heading(section.heading, level=1)
        if section.content:
            doc.add_paragraph(section.content)
    for line in report_block.splitlines():
        doc.add_paragraph(line)
    doc.save(docx_path)
